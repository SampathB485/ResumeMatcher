import streamlit as st
from resume_parser import extract_text_from_resume
from match_engine import compute_match_percentage
from gpt_writer import rewrite_resume
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import re
from jd_parser import * #noqa
from recruiter_tools import (
    generate_recruiter_message,
    generate_cold_email,
    suggest_contact_titles,
    estimate_salary
)
from interview_questions import interview_questions

st.info(
    "Live demo runs on a shared API key with limited daily calls. "
    "If AI features stop responding, clone the repo and run locally "
    "with your own free Gemini API key."
)

st.set_page_config(page_title="Resume Editor", layout="wide")
st.title("JobMatch - Resume & Job Description Analyzer")


# ---- Export Helpers ----

def _parse_markdown_line(line):
    """Determine the type and content of a markdown line."""
    stripped = line.strip()
    
    if stripped.startswith("### "):
        return "h3", stripped[4:].strip()
    elif stripped.startswith("## "):
        return "h2", stripped[3:].strip()
    elif stripped.startswith("# "):
        return "h1", stripped[2:].strip()
    elif stripped.startswith("- ") or stripped.startswith("* "):
        return "bullet", stripped[2:].strip()
    elif stripped == "---" or stripped == "***":
        return "separator", ""
    elif stripped == "":
        return "blank", ""
    else:
        return "text", stripped


def _add_bold_runs(paragraph, text):
    """Parse **bold** markers and add runs with appropriate formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def download_resume_docx(resume_text):
    """Convert markdown-formatted resume text into a properly formatted Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set narrow margins for resume
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    lines = resume_text.split('\n')
    
    for line in lines:
        line_type, content = _parse_markdown_line(line)
        
        if line_type == "blank":
            continue  # Skip blank lines to keep resume tight
            
        elif line_type == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(16)
                run.bold = True
            p.space_after = Pt(4)
            p.space_before = Pt(0)
            
        elif line_type == "h2":
            p = doc.add_paragraph()
            _add_bold_runs(p, content.upper())
            for run in p.runs:
                run.font.size = Pt(12)
                run.bold = True
            p.space_after = Pt(2)
            p.space_before = Pt(8)
            # Add a bottom border to section headers
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): '000000'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif line_type == "h3":
            p = doc.add_paragraph()
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(11)
                run.bold = True
            p.space_after = Pt(1)
            p.space_before = Pt(4)

        elif line_type == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(1)
            p.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            
        elif line_type == "separator":
            continue  # Skip markdown separators
            
        else:  # regular text
            p = doc.add_paragraph()
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(2)
            p.space_before = Pt(0)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def download_resume_pdf(resume_text):
    """Convert markdown-formatted resume text into a properly formatted PDF."""
    buffer = BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles for resume sections
    styles.add(ParagraphStyle(
        name='ResumeName',
        parent=styles['Title'],
        fontSize=16,
        leading=20,
        spaceAfter=4,
        spaceBefore=0,
        alignment=1  # center
    ))
    
    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        leading=16,
        spaceAfter=4,
        spaceBefore=10,
        textColor='black',
        borderWidth=0.5,
        borderPadding=2,
        borderColor='black',
    ))
    
    styles.add(ParagraphStyle(
        name='SubHeader',
        parent=styles['Heading3'],
        fontSize=11,
        leading=14,
        spaceAfter=2,
        spaceBefore=6,
        textColor='black',
    ))
    
    styles.add(ParagraphStyle(
        name='BulletItem',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        spaceAfter=2,
        spaceBefore=0,
        leftIndent=18,
        bulletIndent=6,
    ))
    
    styles.add(ParagraphStyle(
        name='ResumeText',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        spaceAfter=2,
        spaceBefore=0,
    ))

    story = []
    lines = resume_text.split('\n')
    
    for line in lines:
        line_type, content = _parse_markdown_line(line)
        
        # Convert markdown bold to reportlab bold tags
        content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
        
        if line_type == "blank":
            continue
            
        elif line_type == "h1":
            story.append(Paragraph(content, styles['ResumeName']))
            
        elif line_type == "h2":
            # Add a thin line above section headers
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<u><b>{content.upper()}</b></u>", styles['SectionHeader']))
            
        elif line_type == "h3":
            story.append(Paragraph(f"<b>{content}</b>", styles['SubHeader']))
            
        elif line_type == "bullet":
            story.append(Paragraph(f"&bull;  {content}", styles['BulletItem']))
            
        elif line_type == "separator":
            continue
            
        else:
            story.append(Paragraph(content, styles['ResumeText']))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ----- Resume Upload or Paste -----
st.header("Step 1: Upload or Paste Your Resume")

resume_input_method = st.radio("How do you want to provide your resume?", ("Upload File", "Paste Text"))

resume_text = ""
if resume_input_method == "Upload File":
    resume_file = st.file_uploader("Upload your resume (.pdf or .docx)", type=["pdf", "docx"])
    
    if resume_file:
        resume_text = extract_text_from_resume(resume_file).strip()
        st.success("Resume loaded successfully!")

elif resume_input_method == "Paste Text":
    resume_text = st.text_area("Paste your resume here", height=200)

# ----- Job Description Input -----
st.header("Step 2: Paste Job Description")
job_description = st.text_area("Paste the job description here", height=200)

if job_description.strip():
    # Clear cached job info if JD text changed so extraction re-runs
    if "last_jd_hash" not in st.session_state or st.session_state.last_jd_hash != hash(job_description):
        st.session_state.last_jd_hash = hash(job_description)
        st.session_state.jd_info_extracted = False

    job_description_parsed = parse_jd(job_description)
    job_description = job_description_parsed["clean_text"]
    st.success("Job Description parsed successfully!")
else:
    job_description = ""

# ----- Matching & Editing the Resume -----
st.header("Step 3: Match Results")

if resume_text and job_description:
    if st.button("Compute Match %"):
        with st.spinner("Analyzing..."):
            match_percent = compute_match_percentage(resume_text, job_description)
        st.success(f"Resume Match: {match_percent}%")

    if st.button("Rewrite Resume"):
        with st.spinner("Generating optimized resume (this may take a moment for longer resumes)..."):
            edited_resume = rewrite_resume(resume_text, job_description)
        
        # Store in session state so it persists across reruns
        st.session_state["edited_resume"] = edited_resume
        st.success("Edited Resume Ready!")

    # Show download buttons if edited resume exists in session
    if "edited_resume" in st.session_state:
        col1, col2 = st.columns(2)
        
        with col1:
            docx_file = download_resume_docx(st.session_state["edited_resume"])
            st.download_button(
                label="Download Edited Resume (.docx)",
                data=docx_file,
                file_name="Edited_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            pdf_file = download_resume_pdf(st.session_state["edited_resume"])
            st.download_button(
                label="Download Edited Resume (.pdf)",
                data=pdf_file,
                file_name="Edited_Resume.pdf",
                mime="application/pdf"
            )

else:
    st.info("Please upload a resume and paste a job description above to generate results.")

# ----- Reach Out Section - Tools -----
st.header("Step 4: Recruiter Communication & Salary Insights")

if resume_text and job_description:
    if st.button("Generate Recruiter Message"):
        with st.spinner("Creating recruiter message..."):
            recruiter_msg = generate_recruiter_message(job_description)
        st.text_area("LinkedIn Message", recruiter_msg, height=150)

    if st.button("Generate Cold Email"):
        with st.spinner("Writing email..."):
            cold_email = generate_cold_email(job_description)
        st.text_area("Cold Email Template", cold_email, height=450)

    if st.button("Suggest People to Contact"):
        with st.spinner("Finding roles..."):
            titles = (", ".join(suggest_contact_titles(job_description)))
        st.text_area("Recommended Contacts", titles, height=250)

    if st.button("Estimate Salary Range"):
        with st.spinner("Estimating salary..."):
            salary = estimate_salary(job_description)
        st.success(f"Estimated Salary: {salary}")

else:
    st.info("Please upload a resume and paste a job description above to generate results.")


# ----- Potential Interview Questions -----
st.header("Step 5: Potential Interview Prep")

if resume_text and job_description:
    if st.button("Generate Potential Interview Questions"):
        with st.spinner("Generating interview questions..."):
            questions = interview_questions(resume_text, job_description)
        st.text_area("Interview Questions", questions, height=650)

else:
    st.info("Please upload a resume and paste a job description above to generate results.")